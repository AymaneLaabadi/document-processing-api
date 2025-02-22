import logging
import re
from pathlib import Path
from typing import Dict
from docx import Document
from src.config import Config

logger = logging.getLogger(__name__)


class DocumentFiller:
    """Class for filling Word documents with extracted information."""

    def __init__(self, template_path: Path):
        """
        Initialize DocumentFiller with template path.

        Args:
            template_path (Path): Path to the Word template document
        """
        self.template_path = template_path
        try:
            self.document = Document(template_path)
            logger.info(f"Successfully loaded template document: {template_path}")
        except Exception as e:
            logger.error(f"Failed to load template document: {str(e)}")
            raise

    def _replace_placeholders_in_paragraph(self, paragraph, data_dict: Dict) -> None:
        """
        Replace placeholders in a paragraph while preserving formatting.

        Args:
            paragraph: Document paragraph object
            data_dict (Dict): Dictionary containing replacement values
        """
        try:
            full_text = "".join(run.text for run in paragraph.runs)
            modified = False

            for key, value in data_dict.items():
                placeholder = f'#{key}#'
                if re.search(placeholder, full_text, flags=re.IGNORECASE):
                    full_text = re.sub(placeholder, str(value), full_text, flags=re.IGNORECASE)
                    modified = True

            if modified:
                # Store and apply formatting
                remaining_text = full_text
                last_run = paragraph.runs[-1] if paragraph.runs else None

                # Clear existing runs
                for run in paragraph.runs:
                    run.text = ""

                # Create new run with preserved formatting
                if last_run and remaining_text:
                    new_run = paragraph.add_run(remaining_text)
                    new_run.bold = last_run.bold
                    new_run.italic = last_run.italic
                    new_run.underline = last_run.underline
                    new_run.font.size = last_run.font.size
                    new_run.font.name = last_run.font.name
                    if last_run.font.color.rgb:
                        new_run.font.color.rgb = last_run.font.color.rgb

        except Exception as e:
            logger.error(f"Error replacing placeholders in paragraph: {str(e)}")

    def fill_document(self, data_dict: Dict) -> None:
        """
        Fill the document with provided data.

        Args:
            data_dict (Dict): Dictionary containing replacement values
        """
        try:
            # Replace in paragraphs
            for para in self.document.paragraphs:
                self._replace_placeholders_in_paragraph(para, data_dict)

            # Replace in tables
            for table in self.document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            self._replace_placeholders_in_paragraph(para, data_dict)

            logger.info("Successfully filled document with provided data")

        except Exception as e:
            logger.error(f"Error filling document: {str(e)}")
            raise

    def save_document(self, output_path: Path, data_dict: Dict) -> None:
        """
        Save the filled document with a dynamic name based on the "21" field.

        Args:
            output_path (Path): Base path where the document should be saved.
            data_dict (Dict): Dictionary containing extracted values, including "21" for naming.
        """
        try:
            # Convert output_path to a Path object
            output_path = Path(output_path)

            doc_number = data_dict.get("21", "UNKNOWN").replace("/", "_").replace("\\", "_").strip()
            if not doc_number:
                doc_number = "NO_REFERENCE"

            # Ensure the base output filename from config exists
            base_filename = output_path.stem  # Get filename without extension
            file_extension = output_path.suffix  # Get file extension

            # Create new filename by appending "21" value
            new_filename = f"{base_filename}_{doc_number}{file_extension}"
            final_output_path = output_path.parent / new_filename  # Save in same directory

            # Save the document
            self.document.save(final_output_path)
            logger.info(f"Successfully saved filled document to: {final_output_path}")

        except Exception as e:
            logger.error(f"Error saving document: {str(e)}")
            raise
